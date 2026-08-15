from fastapi import FastAPI
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer,Table,TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from io import BytesIO
app=FastAPI(title='Design Mantra Document Service')
class Item(BaseModel):
    areaType:str=''; name:str; vendor:str=''; description:str=''; unit:str='Nos.'; qty:float=1; rate:float=0
class Quote(BaseModel):
    quoteId:str|None=None; invoiceNo:str; client:str; phone:str|None=None; projectType:str|None=None; location:str|None=None; area:float|None=None; validity:str='15 days'; terms:str|None=None; totals:dict; items:list[Item]
@app.get('/health')
def health(): return {'ok':True,'service':'document-service'}
def pdf(q:Quote):
    b=BytesIO(); doc=SimpleDocTemplate(b,pagesize=A4,rightMargin=28,leftMargin=28,topMargin=28,bottomMargin=28); s=getSampleStyleSheet(); story=[Paragraph('DESIGN MANTRA',s['Title']),Paragraph(f'Estimate {q.invoiceNo}',s['Normal']),Spacer(1,12),Paragraph(f'<b>Client:</b> {q.client} &nbsp;&nbsp; <b>Project:</b> {q.projectType or ""} &nbsp;&nbsp; <b>Location:</b> {q.location or ""}',s['Normal']),Spacer(1,12)]
    data=[['#','Area','Item','Vendor','Unit','Qty','Rate','Amount']]
    for i,x in enumerate(q.items,1): data.append([i,x.areaType,x.name,x.vendor,x.unit,x.qty,x.rate,x.qty*x.rate])
    t=Table(data,repeatRows=1); t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#f5eee6')),('TEXTCOLOR',(0,0),(-1,0),colors.HexColor('#9a730d')),('GRID',(0,0),(-1,-1),.35,colors.HexColor('#d6b67c')),('ALIGN',(5,1),(-1,-1),'RIGHT')])) ;story += [t,Spacer(1,14)]
    z=q.totals; story += [Paragraph(f'Subtotal: ₹ {z.get("subtotal",0):,.2f}',s['Normal']),Paragraph(f'Discount: ₹ {z.get("discount",0):,.2f}',s['Normal']),Paragraph(f'GST: ₹ {z.get("tax",0):,.2f}',s['Normal']),Paragraph(f'<b>Grand Total: ₹ {z.get("total",0):,.2f}</b>',s['Heading2']),Spacer(1,10),Paragraph(q.terms or 'Terms & conditions as agreed with client.',s['Normal'])]; doc.build(story); b.seek(0); return b
@app.post('/generate/pdf')
def generate_pdf(q:Quote): return StreamingResponse(pdf(q),media_type='application/pdf',headers={'Content-Disposition':f'attachment; filename={q.invoiceNo}.pdf'})
@app.post('/generate/docx')
def generate_docx(q:Quote):
    b=BytesIO(); d=Document(); d.add_heading('DESIGN MANTRA',0); d.add_paragraph(f'Estimate: {q.invoiceNo}\nClient: {q.client}\nProject: {q.projectType or ""}\nLocation: {q.location or ""}'); table=d.add_table(rows=1,cols=8); hdr=table.rows[0].cells
    for c,v in zip(hdr,['#','Area','Item','Vendor','Unit','Qty','Rate','Amount']): c.text=v
    for i,x in enumerate(q.items,1):
        c=table.add_row().cells
        for cell,v in zip(c,[i,x.areaType,x.name,x.vendor,x.unit,x.qty,x.rate,x.qty*x.rate]): cell.text=str(v)
    z=q.totals; d.add_paragraph(f'Subtotal: ₹ {z.get("subtotal",0):,.2f}\nDiscount: ₹ {z.get("discount",0):,.2f}\nGST: ₹ {z.get("tax",0):,.2f}\nGrand Total: ₹ {z.get("total",0):,.2f}'); d.add_paragraph(q.terms or 'Terms & conditions as agreed with client.'); d.save(b); b.seek(0); return StreamingResponse(b,media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',headers={'Content-Disposition':f'attachment; filename={q.invoiceNo}.docx'})
